"""
OmniFind Web 服务 —— FastAPI,监听 127.0.0.1,提供查询 API + 静态 UI。
"""
from __future__ import annotations
from pathlib import Path
from contextlib import asynccontextmanager
import threading
import time

from fastapi import FastAPI, Query
from fastapi.responses import FileResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
import logging
import os

from omnifind.core.config import OmniConfig, ensure_dirs
from omnifind.core.router import QueryRouter
from omnifind.layers.l1_filename.index import FilenameIndex
from omnifind.layers.l2_fulltext.index import FullTextIndex

STATIC_DIR = Path(__file__).resolve().parent / "static"

# 全局单例(简单起见;后续可换依赖注入)
_state: dict = {}

# 索引任务状态
_index_task: dict = {
    "running": False,
    "type": "",      # l1 / l2 / all
    "progress": 0,
    "total": 0,
    "message": "",
    "start_time": 0,
    "error": "",
}


@asynccontextmanager
async def lifespan(app: FastAPI):
    ensure_dirs()
    cfg = OmniConfig.load()
    l1 = FilenameIndex()
    l2 = FullTextIndex()
    l3 = None
    try:
        if cfg.resolved_embed_model_path.joinpath("model.onnx").exists():
            from omnifind.layers.l3_semantic.builder import make_embedder
            from omnifind.layers.l3_semantic.index import SemanticIndex
            emb = make_embedder(cfg)
            l3 = SemanticIndex(emb, dim=emb.dim, min_score=cfg.semantic_min_score)
    except Exception as e:  # noqa: BLE001
        print(f"[L3] 语义层不可用(跳过):{e}")
    router = QueryRouter(l1=l1, l2=l2, l3=l3)
    _state.update(cfg=cfg, l1=l1, l2=l2, l3=l3, router=router)

    # 首次自动构建:新机器/空索引时,启动后台构建(不阻塞启动,离线可跑——模型随包走)。
    # L1/L2 空 -> 触发全量;L3 可用且已配置 semantic_dirs 且空 -> 触发语义构建。
    try:
        if l1.count() == 0:
            start_build("all")          # 含 L3(当且仅当已配置 semantic_dirs)
        elif l2.count() == 0:
            start_build("l2")
        # 独立 L3 首建:L1/L2 已存在但语义库为空
        if l3 is not None and cfg.semantic_dirs and l3.count() == 0:
            start_build("l3")
    except Exception as e:  # noqa: BLE001
        print(f"[auto-build] 触发失败: {e}")

    # 启动 L3 增量监听:日常增删改只更新单文件向量,无需全量重建语义索引
    if l3 is not None and cfg.semantic_dirs:
        try:
            from omnifind.layers.l3_semantic.watcher import SemanticWatcher
            watcher = SemanticWatcher(cfg, l3, interval=30.0)
            watcher.start()
            _state["l3_watcher"] = watcher
        except Exception as e:  # noqa: BLE001
            print(f"[L3 watch] 启动失败(跳过): {e}")

    yield
    watcher = _state.get("l3_watcher")
    if watcher:
        watcher.stop()
    l1.close()
    l2.close()


app = FastAPI(title="OmniFind", version="0.2.0", lifespan=lifespan)


logger = logging.getLogger(__name__)


@app.get("/api/search")
def search(q: str = Query("", max_length=2000),
           limit: int = Query(30, ge=1, le=5000),
           offset: int = Query(0, ge=0, le=100000),
           mode: str | None = Query(None, pattern="^(auto|filename|fulltext|semantic|all|filename_regex)$"),
           ext: str | None = None,
           sort: str | None = None):
    router: QueryRouter = _state["router"]
    # ext 统一补点,防传 py 命中 happy
    if ext and not ext.startswith("."):
        ext = "." + ext
    try:
        # 多取 1 条用于精确判断 has_more
        resp = router.search(q, limit=limit + offset + 1, mode=mode, ext=ext, sort=sort)
        total_got = len(resp.hits)
        # 分页切片:offset 越界必须返回空,不能回落到第一页
        if offset > 0:
            resp.hits = resp.hits[offset:] if total_got > offset else []
        has_more = len(resp.hits) > limit
        resp.hits = resp.hits[:limit]
        result = resp.to_dict()
        result["has_more"] = has_more
        return JSONResponse(result)
    except Exception as e:
        # 详细错误只写日志, 不向客户端回吐原始数据库/系统错误串(避免信息泄漏与 XSS 注入)
        logger.exception("搜索异常 query=%s mode=%s", q, mode)
        return JSONResponse({
            "query": q,
            "mode": "error",
            "hits": [],
            "counts": {"error": "搜索失败，请简化关键词或稍后重试"},
            "has_more": False,
        }, status_code=500)


@app.get("/api/status")
def status():
    l3 = _state.get("l3")
    cfg = _state["cfg"]
    return {
        "l1_count": _state["l1"].count(),
        "l2_count": _state["l2"].count(),
        "l3_count": l3.count() if l3 else 0,
        "scan_roots": cfg.scan_roots,
        "exclude_dirs": cfg.exclude_dirs,
        "fulltext_exts": cfg.fulltext_exts,
        "max_fulltext_mb": cfg.max_fulltext_mb,
        "host": cfg.host,
        "port": cfg.port,
        "data_dir": str(cfg.data_dir_path),
        "l3_available": l3 is not None,
    }


@app.get("/api/config")
def get_config():
    cfg = _state["cfg"]
    return {
        "scan_roots": cfg.scan_roots,
        "exclude_dirs": cfg.exclude_dirs,
        "fulltext_exts": cfg.fulltext_exts,
        "max_fulltext_mb": cfg.max_fulltext_mb,
        "l1_backend": cfg.l1_backend,
        "semantic_dirs": cfg.semantic_dirs,
        "semantic_exts": cfg.semantic_exts,
        "semantic_full_disk": cfg.semantic_full_disk,
        "host": cfg.host,
        "port": cfg.port,
        "data_dir": str(cfg.data_dir_path),
    }


# 允许通过 API 修改的字段白名单及类型（防任意字段注入导致运行态不一致）
_CONFIG_SCHEMA: dict[str, type] = {
    "scan_roots": list,
    "exclude_dirs": list,
    "fulltext_exts": list,
    "semantic_dirs": list,
    "semantic_exts": list,
    "semantic_full_disk": bool,
    "max_fulltext_mb": int,
    "chunk_size": int,
    "chunk_overlap": int,
    "semantic_min_score": float,
    "l1_backend": str,
}


@app.post("/api/config")
def update_config(data: dict):
    cfg = _state["cfg"]
    # 类型校验：字段必须在白名单内且类型匹配
    rejected: dict[str, str] = {}
    for k, v in data.items():
        expected = _CONFIG_SCHEMA.get(k)
        if expected is None:
            rejected[k] = "字段不允许在线修改"
        elif v is not None and not isinstance(v, expected):
            rejected[k] = f"类型应为 {expected.__name__}"
        elif expected is list and v is not None and not all(isinstance(i, str) for i in v):
            rejected[k] = "列表元素必须为字符串"
    if rejected:
        return JSONResponse({"ok": False, "error": f"非法字段: {rejected}"}, status_code=400)

    changed = False
    for k, v in data.items():
        if v is not None:
            setattr(cfg, k, v)
            # 阈值类运行态参数需同步到已实例化的对象,否则热更新不生效
            if k == "semantic_min_score" and _state.get("l3") is not None:
                _state["l3"].min_score = float(v)
            changed = True
    if changed:
        # 保存到 config.local.yaml
        from pathlib import Path as _P
        import yaml
        ROOT = _P(__file__).resolve().parent.parent.parent
        local_cfg_path = ROOT / "config.local.yaml"
        try:
            existing = {}
            if local_cfg_path.exists():
                existing = yaml.safe_load(local_cfg_path.read_text(encoding="utf-8")) or {}
            existing.update(data)
            local_cfg_path.write_text(
                yaml.dump(existing, allow_unicode=True, default_flow_style=False),
                encoding="utf-8"
            )
            return {"ok": True, "saved": True, "path": str(local_cfg_path)}
        except Exception as e:
            logger.exception("保存配置失败")
            return JSONResponse({"ok": False, "error": str(e)}, status_code=500)
    return {"ok": True, "saved": False}


@app.get("/api/index/status")
def index_status():
    return {
        "running": _index_task["running"],
        "type": _index_task["type"],
        "progress": _index_task["progress"],
        "total": _index_task["total"],
        "message": _index_task["message"],
        "start_time": _index_task["start_time"],
        "elapsed": time.time() - _index_task["start_time"] if _index_task["running"] else 0,
        "error": _index_task["error"],
    }


_rebuild_lock = threading.Lock()


def _run_build(type: str) -> None:
    """实际构建逻辑(后台线程调用)。支持 l1/l2/l3/all。"""
    cfg = _state["cfg"]
    l1 = _state["l1"]
    l2 = _state["l2"]
    l3 = _state.get("l3")
    try:
        if type in ("l1", "all"):
            _index_task["message"] = "正在重建文件名索引..."
            from omnifind.core.indexer import build_filename_index
            l1.clear()
            n = build_filename_index(cfg, l1)
            _index_task["progress"] = n
            _index_task["total"] = n
            _index_task["message"] = f"文件名索引完成: {n} 条"

        if type in ("l2", "all"):
            _index_task["message"] = "正在重建全文索引..."
            from omnifind.core.indexer import build_fulltext_index
            l2.clear()
            n = build_fulltext_index(cfg, l2)
            _index_task["progress"] = n
            _index_task["total"] = n
            _index_task["message"] = f"全文索引完成: {n} 篇"

        if type in ("l3", "all"):
            if l3 is None:
                raise RuntimeError("语义层未启用(模型/依赖缺失)，无法重建 L3")
            # 安全阀:未配置语义目录且未明确要求全盘,绝不偷偷向量化全盘
            if not cfg.semantic_dirs and not cfg.semantic_full_disk:
                logger.warning(
                    "[L3] 未配置 semantic_dirs 且 semantic_full_disk=False，跳过语义构建(避免全盘向量化)"
                )
            else:
                _index_task["message"] = "正在重建语义索引..."
                from omnifind.layers.l3_semantic.builder import build_semantic_index
                l3.drop()
                build_semantic_index(cfg, l3)
                _index_task["message"] = f"语义索引完成: {l3.count()} chunks"

        _index_task["message"] += " · 全部完成"
    except Exception as e:
        logger.exception("索引重建失败")
        _index_task["error"] = str(e)
        _index_task["message"] = f"失败: {e}"
    finally:
        _index_task["running"] = False


def start_build(type: str) -> bool:
    """原子启动一次构建任务;若已有任务在进行中返回 False。"""
    with _rebuild_lock:
        if _index_task["running"]:
            return False
        _index_task["running"] = True
    _index_task["type"] = type
    _index_task["progress"] = 0
    _index_task["total"] = 0
    _index_task["message"] = "准备中..."
    _index_task["start_time"] = time.time()
    _index_task["error"] = ""
    t = threading.Thread(target=_run_build, args=(type,), daemon=True)
    t.start()
    return True


@app.post("/api/index/rebuild")
def rebuild_index(type: str = Query("all", pattern="^(l1|l2|l3|all)$")):
    ok = start_build(type)
    if not ok:
        return JSONResponse({"ok": False, "error": "索引任务正在进行中，请稍候"}, status_code=400)
    return {"ok": True, "started": True, "type": type}


@app.get("/")
def index():
    return FileResponse(STATIC_DIR / "index.html")


# 支持预览的纯文本扩展名
_TEXT_EXTS = {
    ".txt", ".md", ".py", ".js", ".ts", ".java", ".go", ".rs",
    ".c", ".cpp", ".h", ".hpp", ".cs", ".rb", ".php", ".swift",
    ".kt", ".scala", ".r", ".lua", ".perl", ".sh", ".bat",
    ".ps1", ".cmd", ".bash", ".zsh", ".fish", ".yaml", ".yml",
    ".json", ".xml", ".html", ".htm", ".css", ".scss", ".less",
    ".sql", ".log", ".csv", ".ini", ".toml", ".cfg", ".conf",
    ".env", ".gitignore", ".dockerfile", ".makefile", ".cmake",
    ".proto", ".graphql", ".vue", ".svelte", ".jsx", ".tsx",
    ".tcl", ".ex", ".exs", ".erl", ".hs", ".ml", ".mli",
    ".vim", ".el", ".lisp", ".clj", ".cljs", ".rkt",
    ".dockerignore", ".editorconfig", ".prettierrc", ".eslintrc",
    "Dockerfile", "Makefile", "README", "LICENSE", "CHANGELOG",
    "CONTRIBUTING", ".gitattributes", ".mailmap",
}
# 需要用专用库提取文本的文档扩展名
_DOC_EXTS = {".docx", ".xlsx", ".xls"}
# 预览最大 500KB（增大限制，用户反馈预览内容不全）
_MAX_PREVIEW_SIZE = 500 * 1024


@app.get("/api/preview")
def preview(path: str = Query(...)):
    """读取文件内容用于预览。优先从 L2 索引取 body，否则直接读文件。"""
    ok, resolved, err = _verify_path_in_index(path)
    if not ok:
        status = 400 if "必填" in (err or "") or "无效" in (err or "") else (404 if "不存在" in (err or "") else 403)
        return JSONResponse({"ok": False, "error": err}, status_code=status)

    if resolved.is_dir():
        return JSONResponse({"ok": False, "error": "是目录，无法预览"}, status_code=400)

    p = resolved
    ext = p.suffix.lower()
    name = p.name

    # OOM 预检：任何分支动手前先 stat 大小，超限直接拒绝（不整读进内存）
    try:
        fsize = p.stat().st_size
    except OSError as e:
        return JSONResponse({"ok": False, "error": f"无法读取文件属性: {e}"}, status_code=500)
    _HARD_LIMIT = _MAX_PREVIEW_SIZE * 100  # 50MB：文档/PDF 提取库也会整读，统一硬顶
    if fsize > _HARD_LIMIT:
        return JSONResponse({
            "ok": True, "path": path, "content": None,
            "source": "too_large",
            "message": f"文件过大({fsize // (1024 * 1024)}MB)，超过预览硬限制(50MB)。",
        })

    # ===== 核心策略：预览优先读原始文件，L2 body 仅作最后 fallback =====
    # 原因: L2 存的是 jieba 分词后的 token(空格分隔)，格式全丢，用户看不懂

    # 1) 文档格式 (.docx/.xlsx/.xls) → 用专用库提取
    if ext in _DOC_EXTS:
        doc_result = _extract_document_text(p, ext)
        if doc_result is not None:
            content, src_label = doc_result
            truncated = len(content) > _MAX_PREVIEW_SIZE
            return JSONResponse({
                "ok": True, "path": path,
                "content": content[:_MAX_PREVIEW_SIZE],
                "source": src_label, "size": p.stat().st_size,
                "truncated": truncated,
            })

    # 2) PDF → fitz / pdfplumber / PyPDF2 提取
    if ext == ".pdf":
        pdf_result = _extract_pdf_text(p)
        if pdf_result is not None:
            content, src_label = pdf_result
            truncated = len(content) > _MAX_PREVIEW_SIZE
            return JSONResponse({
                "ok": True, "path": path,
                "content": content[:_MAX_PREVIEW_SIZE],
                "source": src_label, "size": p.stat().st_size,
                "truncated": truncated,
            })
        else:
            return JSONResponse({
                "ok": True, "path": path, "content": None,
                "source": "pdf_no_lib",
                "message": "PDF 文本提取失败（文件可能为扫描件/图片型 PDF，或解析库均不可用）。",
            })

    # 3) 判断是否为文本/代码文件
    is_text = (ext in _TEXT_EXTS or name in _TEXT_EXTS or _is_text_file(p))

    if not is_text:
        return JSONResponse({
            "ok": True, "path": path, "content": None,
            "source": "binary_skip",
            "message": f"二进制/不支持的格式({ext or '未知类型'})。支持预览: 纯文本/.docx/.xlsx/.xls/.pdf(需pdfplumber)。",
        })

    # 4) 文本文件 → 直接读取原始内容（保留原始格式！）
    try:
        # 先按 stat 大小拒绝，再读 —— 不允许"整读后才发现太大"
        if fsize > _MAX_PREVIEW_SIZE * 10:
            return JSONResponse({
                "ok": True, "path": path, "content": None,
                "source": "too_large",
                "message": f"文件过大({fsize//1024}KB)，超过预览限制。",
            })
        raw = p.read_bytes()
        file_size = len(raw)
        for enc in ("utf-8", "gbk", "gb2312", "latin-1"):
            try:
                text = raw.decode(enc)
                truncated = len(text) > _MAX_PREVIEW_SIZE
                return JSONResponse({
                    "ok": True, "path": path,
                    "content": text[:_MAX_PREVIEW_SIZE],
                    "source": "file_read",
                    "encoding": enc,
                    "size": file_size,
                    "truncated": truncated,
                })
            except (UnicodeDecodeError, LookupError):
                continue
        return JSONResponse({"ok": False, "error": "无法以已知编码解码文件"}, status_code=422)
    except Exception as e:
        logger.exception("文件读取失败 path=%s: %s", path, e)
        return JSONResponse({"ok": False, "error": f"读取失败: {e}"}, status_code=500)


def _is_text_file(p: Path) -> bool:
    """启发式检测：读一小段判断是否为文本文件。"""
    try:
        raw = p.read_bytes()
        chunk = raw[:min(8192, len(raw))]
        if b'\x00' in chunk:
            return False
        chunk.decode('utf-8')
        return True
    except (OSError, UnicodeDecodeError):
        return False


def _extract_document_text(p: Path, ext: str) -> tuple[str, str] | None:
    """
    从 Office 文档中提取纯文本内容。
    返回 (text, source_label) 或 None（解析失败）。
    """
    try:
        if ext == ".docx":
            import docx
            doc = docx.Document(str(p))
            parts = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    parts.append(text)
            # 也尝试提取表格中的文字
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        parts.append(" | ".join(row_text))
            body = "\n\n".join(parts)
            return (body, "docx_extract") if body else None

        elif ext == ".xlsx":
            import openpyxl
            wb = openpyxl.load_workbook(str(p), read_only=True, data_only=True)
            parts = []
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                parts.append(f"=== {sheet_name} ===")
                for row in sheet.iter_rows(values_only=True):
                    # 过滤全空行
                    row_text = [str(c) if c is not None else "" for c in row]
                    if any(cell.strip() for cell in row_text if cell):
                        parts.append("\t".join(row_text))
            wb.close()
            body = "\n".join(parts)
            return (body, "xlsx_extract") if body else None

        elif ext == ".xls":
            import xlrd
            wb = xlrd.open_workbook(str(p))
            parts = []
            for sheet in wb.sheets():
                parts.append(f"=== {sheet.name} ===")
                for row_idx in range(sheet.nrows):
                    row = sheet.row_values(row_idx)
                    row_text = [str(c) if c is not None else "" for c in row]
                    if any(cell.strip() for cell in row_text if cell):
                        parts.append("\t".join(row_text))
            body = "\n".join(parts)
            return (body, "xls_extract") if body else None

    except ImportError as e:
        logger.debug("文档库缺失 ext=%s: %s", ext, e)
        return None
    except Exception as e:
        logger.debug("文档解析失败 path=%s ext=%s: %s", p, ext, e)
        return None


def _extract_pdf_text(p: Path) -> tuple[str, str] | None:
    """从 PDF 中提取文本。优先 PyMuPDF/fitz（项目已声明依赖），
    其次 pdfplumber，最后 PyPDF2。"""
    # 优先 fitz（requirements.txt 已声明 pymupdf，与索引抽取器一致）
    try:
        import fitz  # PyMuPDF
        parts = []
        with fitz.open(str(p)) as doc:
            for i, page in enumerate(doc):
                text = page.get_text("text")
                if text and text.strip():
                    parts.append(f"--- Page {i+1} ---\n{text}")
        body = "\n\n".join(parts)
        return (body, "pdf_fitz") if body else None
    except ImportError:
        pass
    except Exception as e:
        logger.debug("fitz 解析失败: %s", e)

    # fallback: pdfplumber（更好的表格/布局支持）
    try:
        import pdfplumber
        parts = []
        with pdfplumber.open(str(p)) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text()
                if text and text.strip():
                    parts.append(f"--- Page {i+1} ---\n{text}")
        body = "\n\n".join(parts)
        return (body, "pdf_plumber") if body else None
    except ImportError:
        pass
    except Exception as e:
        logger.debug("pdfplumber 解析失败: %s", e)

    # fallback: PyPDF2
    try:
        import PyPDF2
        reader = PyPDF2.PdfReader(str(p))
        parts = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text()
            if text and text.strip():
                parts.append(f"--- Page {i+1} ---\n{text}")
        body = "\n\n".join(parts)
        return (body, "pdf_pypdf2") if body else None
    except ImportError:
        pass
    except Exception as e:
        logger.debug("PyPDF2 解析失败: %s", e)

    return None


def _verify_path_in_index(path_str: str) -> tuple[bool, Path | None, str | None]:
    """验证文件路径是否在索引中，返回 (是否在索引中, 解析后的路径, 错误信息)。"""
    if not path_str:
        return False, None, "path 参数必填"
    try:
        resolved = Path(path_str).resolve()
    except (OSError, ValueError) as e:
        return False, None, f"无效路径: {e}"
    l1 = _state.get("l1")
    l2 = _state.get("l2")
    in_index = False
    if l2:
        try:
            row = l2.conn.execute(
                "SELECT 1 FROM files WHERE path=? LIMIT 1", (str(resolved),)
            ).fetchone()
            if row:
                in_index = True
        except Exception:
            pass
    if not in_index and l1:
        try:
            row = l1.conn.execute(
                "SELECT 1 FROM entries WHERE path=? LIMIT 1", (str(resolved),)
            ).fetchone()
            if row:
                in_index = True
        except Exception:
            pass
    if not in_index:
        return False, resolved, "文件不在索引中，无法操作（安全限制：仅可操作已索引文件）"
    if not resolved.exists():
        return False, resolved, "文件不存在"
    return True, resolved, None


@app.post("/api/file/open")
def open_file(path: str = Query(...)):
    """打开文件（用系统默认程序）。"""
    ok, resolved, err = _verify_path_in_index(path)
    if not ok:
        return JSONResponse({"ok": False, "error": err}, status_code=400 if "必填" in err or "无效" in err else 403)
    try:
        import subprocess, platform
        if platform.system() == "Windows":
            os.startfile(str(resolved))
        elif platform.system() == "Darwin":
            subprocess.run(["open", str(resolved)])
        else:
            subprocess.run(["xdg-open", str(resolved)])
        return JSONResponse({"ok": True, "path": str(resolved)})
    except Exception as e:
        logger.exception("打开文件失败 path=%s", path)
        return JSONResponse({"ok": False, "error": f"打开失败: {e}"}, status_code=500)


@app.post("/api/file/reveal")
def reveal_file(path: str = Query(...)):
    """在文件管理器中显示文件（打开所在位置并选中）。"""
    ok, resolved, err = _verify_path_in_index(path)
    if not ok:
        return JSONResponse({"ok": False, "error": err}, status_code=400 if "必填" in err or "无效" in err else 403)
    try:
        import subprocess, platform
        if platform.system() == "Windows":
            subprocess.run(["explorer", "/select,", str(resolved)])
        elif platform.system() == "Darwin":
            subprocess.run(["open", "-R", str(resolved)])
        else:
            parent = resolved.parent
            subprocess.run(["xdg-open", str(parent)])
        return JSONResponse({"ok": True, "path": str(resolved)})
    except Exception as e:
        logger.exception("定位文件失败 path=%s", path)
        return JSONResponse({"ok": False, "error": f"定位失败: {e}"}, status_code=500)


@app.get("/api/file/info")
def file_info(path: str = Query(...)):
    """获取文件详细信息（大小、修改时间、类型等）。"""
    ok, resolved, err = _verify_path_in_index(path)
    if not ok:
        return JSONResponse({"ok": False, "error": err}, status_code=400 if "必填" in err or "无效" in err else 403)
    try:
        st = resolved.stat()
        return JSONResponse({
            "ok": True,
            "path": str(resolved),
            "name": resolved.name,
            "ext": resolved.suffix.lower(),
            "size": st.st_size,
            "mtime": st.st_mtime,
            "ctime": st.st_ctime,
            "is_dir": resolved.is_dir(),
        })
    except Exception as e:
        logger.exception("获取文件信息失败 path=%s", path)
        return JSONResponse({"ok": False, "error": f"获取失败: {e}"}, status_code=500)


# 静态资源
if STATIC_DIR.exists():
    app.mount("/static", StaticFiles(directory=str(STATIC_DIR)), name="static")


def main():
    import uvicorn
    cfg = OmniConfig.load()
    uvicorn.run("omnifind.web.server:app", host=cfg.host, port=cfg.port, reload=False)


if __name__ == "__main__":
    main()
